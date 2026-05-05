"""DOCX exporter — converts Markdown into a Word document.

Maps:
    # → Heading 1
    ## → Heading 2
    ### → Heading 3
    -/* list → bulleted list
    1. list → numbered list
    paragraph → body text

Requires `python-docx` (install via `pip install praxia[office]`).
"""
from __future__ import annotations

import re
from io import BytesIO
from typing import Any


class DocxExporter:
    format = "docx"
    extensions = ("docx",)

    def __init__(
        self,
        *,
        title: str | None = None,
        author: str | None = None,
    ) -> None:
        self.title = title
        self.author = author

    def export(self, content: Any) -> bytes:
        try:
            from docx import Document  # type: ignore[import-untyped]
        except ImportError as e:
            raise ImportError(
                "python-docx is required for DOCX export. "
                "Install with: pip install 'praxia[office]'"
            ) from e

        if isinstance(content, dict):
            content = self._dict_to_md(content)
        elif not isinstance(content, str):
            content = str(content)

        doc = Document()
        if self.title:
            doc.add_heading(self.title, level=0)
        if self.author:
            p = doc.add_paragraph()
            run = p.add_run(self.author)
            run.italic = True

        in_code = False
        code_buf: list[str] = []
        for raw in content.splitlines():
            line = raw.rstrip()
            if line.startswith("```"):
                if in_code:
                    doc.add_paragraph("\n".join(code_buf), style="Intense Quote")
                    code_buf = []
                    in_code = False
                else:
                    in_code = True
                continue
            if in_code:
                code_buf.append(line)
                continue

            stripped = line.lstrip()
            m_h = re.match(r"^(#{1,6})\s+(.*)$", stripped)
            if m_h:
                level = min(len(m_h.group(1)), 4)
                doc.add_heading(m_h.group(2), level=level)
                continue
            if re.match(r"^[-*+]\s+", stripped):
                text = re.sub(r"^[-*+]\s+", "", stripped)
                doc.add_paragraph(text, style="List Bullet")
                continue
            if re.match(r"^\d+\.\s+", stripped):
                text = re.sub(r"^\d+\.\s+", "", stripped)
                doc.add_paragraph(text, style="List Number")
                continue
            if not stripped:
                continue
            doc.add_paragraph(stripped)

        if in_code and code_buf:
            doc.add_paragraph("\n".join(code_buf), style="Intense Quote")

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def _dict_to_md(d: dict[str, Any]) -> str:
        from praxia.io.exporters.md_exporter import MarkdownExporter
        return MarkdownExporter._dict_to_md(d)
