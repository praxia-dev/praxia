"""Microsoft Word (.docx) parser — uses python-docx for text + tables
and a bare zipfile read for the embedded `word/media/*` images so the
vision LLM can see charts / diagrams / screenshots that the .docx
contains."""
from __future__ import annotations

import base64
import io
import re
import zipfile
from typing import Any

from praxia.io.parsers.base import ParsedFile

# Caps to keep the metadata payload sane. A pathological .docx with
# 200 high-res images would otherwise inflate every chunk record we
# write to disk and every JSON response we send.
_MAX_PER_IMAGE_BYTES = 5 * 1024 * 1024          # 5 MB per image
_MAX_TOTAL_IMAGE_BYTES = 20 * 1024 * 1024       # 20 MB embedded total
_MEDIA_PATH_RE = re.compile(r"^word/media/.*\.(png|jpe?g|gif|webp)$", re.IGNORECASE)


def _ext_to_mime(name: str) -> str | None:
    n = name.lower()
    if n.endswith(".png"):  return "image/png"
    if n.endswith(".jpg") or n.endswith(".jpeg"): return "image/jpeg"
    if n.endswith(".gif"):  return "image/gif"
    if n.endswith(".webp"): return "image/webp"
    return None


def _extract_embedded_images(data: bytes) -> list[dict[str, Any]]:
    """Return a list of {name, mime, data (base64), bytes_size} from the
    .docx's `word/media/` folder. Drops images larger than the per-image
    cap; stops adding once the cumulative total exceeds the cap (rest
    are summarised by count in the parser's return metadata).
    """
    out: list[dict[str, Any]] = []
    total = 0
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            for info in zf.infolist():
                if not _MEDIA_PATH_RE.match(info.filename):
                    continue
                if info.file_size > _MAX_PER_IMAGE_BYTES:
                    continue
                if total + info.file_size > _MAX_TOTAL_IMAGE_BYTES:
                    break
                mime = _ext_to_mime(info.filename)
                if not mime:
                    continue
                try:
                    raw = zf.read(info.filename)
                except KeyError:
                    continue
                out.append({
                    "name": info.filename.split("/")[-1],
                    "mime": mime,
                    "data": base64.b64encode(raw).decode("ascii"),
                    "bytes_size": info.file_size,
                })
                total += info.file_size
    except zipfile.BadZipFile:
        # The bytes weren't a valid ZIP — python-docx would have raised
        # too, so the caller will see that error first. Return empty so
        # extraction doesn't crash the parse path.
        return []
    return out


class DocxParser:
    name = "docx"

    def parse(self, data: bytes, *, filename: str, **kwargs: Any) -> ParsedFile:
        try:
            import docx
        except ImportError as e:
            raise ImportError(
                "Word (.docx) parsing requires `python-docx`. Install with:\n"
                '  pip install "praxia[office]"'
            ) from e

        doc = docx.Document(io.BytesIO(data))

        # Extract paragraphs with style for structure preservation
        sections: list[tuple[str, str]] = []
        current_heading = "Body"
        body_lines: list[str] = []
        full_text: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = (para.style.name or "").lower() if para.style else ""
            if "heading" in style:
                # Flush previous section
                if body_lines:
                    sections.append((current_heading, "\n".join(body_lines)))
                    body_lines = []
                current_heading = text
                full_text.append(f"\n## {text}\n")
            else:
                body_lines.append(text)
                full_text.append(text)

        # Flush last section
        if body_lines:
            sections.append((current_heading, "\n".join(body_lines)))

        # Extract tables as Markdown
        for i, table in enumerate(doc.tables, start=1):
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows:
                full_text.append(f"\n### Table {i}\n")
                full_text.append("| " + " | ".join(rows[0]) + " |")
                full_text.append("|" + "|".join(["---"] * len(rows[0])) + "|")
                for row in rows[1:]:
                    full_text.append("| " + " | ".join(row) + " |")

        # Embedded images — alpha20 multi-modal pipeline. Vision-capable
        # LLMs read these from chunk metadata at retrieval time.
        embedded = _extract_embedded_images(data)

        return ParsedFile(
            filename=filename,
            content="\n".join(full_text),
            metadata={
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "section_count": len(sections),
                "embedded_images": embedded,
                "embedded_image_count": len(embedded),
            },
            sections=sections,
        )
